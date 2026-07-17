"""Extract structured data from transcript files using the Claude API.

PDFs are sent to Claude as documents (works for both text-based and scanned
transcripts); .docx files are converted to plain text first. Extraction results
are cached on disk keyed by file content hash, so re-runs don't re-pay for
extraction.
"""

import base64
import hashlib
import json
import logging
import re
from pathlib import Path

import anthropic
import pydantic

from .config import DEFAULT_MODEL
from .models import TranscriptData

log = logging.getLogger(__name__)

EXTRACTION_PROMPT = """\
You are parsing a college/university academic transcript.

Extract ALL of the following into the structured output:
- The issuing institution (the school that produced this transcript), its
  location if printed, and its district/system name if it is part of one
  (e.g. a California community college district).
- The date the transcript was issued/printed.
- Every institution mentioned as a source of transfer credit.
- Every degree, certificate, or credential awarded (name, major/concentration,
  award date, honors, and awarding institution if it differs from the issuer).
- EVERY course line on the transcript, including transfer-credit courses,
  test credit (AP/CLEP), and in-progress courses. For each course capture the
  subject code, course number, full title, credits attempted/earned, grade,
  term label, term dates if printed, level (undergraduate/graduate), whether
  it is transfer credit, and which institution taught it if not the issuer.
  Transcripts often use multi-column layouts - read carefully and do not skip
  columns or pages.
- Convert any dates you extract to ISO 8601 (YYYY-MM-DD) when the full date is
  printed; otherwise keep the text as printed (e.g. "Fall 2019").

Do NOT extract any student personally identifiable information: no student
names, student IDs, SSNs, birth dates, or addresses. Leave those out entirely.

If parts of the document are unreadable or ambiguous, note that briefly in
parsing_notes rather than guessing.

Respond with ONLY a single JSON object (no markdown fences, no commentary)
that conforms to this JSON Schema:

{schema}
"""


def _file_hash(path: Path) -> str:
    h = hashlib.sha256()
    h.update(path.read_bytes())
    return h.hexdigest()[:24]


def _docx_to_text(path: Path) -> str:
    import docx  # python-docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(p for p in parts if p.strip())


def _prompt() -> str:
    schema = json.dumps(TranscriptData.model_json_schema(), indent=1)
    return EXTRACTION_PROMPT.format(schema=schema)


def build_content(path: Path) -> list[dict]:
    """Build the user-message content blocks for one transcript file."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        data = base64.standard_b64encode(path.read_bytes()).decode("ascii")
        return [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": data,
                },
            },
            {"type": "text", "text": _prompt()},
        ]
    if suffix == ".docx":
        text = _docx_to_text(path)
        return [
            {"type": "text", "text": f"<transcript_text>\n{text}\n</transcript_text>"},
            {"type": "text", "text": _prompt()},
        ]
    raise ValueError(f"Unsupported file type: {path.name}")


def _parse_json_response(raw: str) -> TranscriptData:
    """Parse the model's JSON reply, tolerating markdown fences."""
    text = raw.strip()
    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", text, re.S)
    if fence:
        text = fence.group(1)
    # If any prose slipped in, isolate the outermost JSON object.
    if not text.startswith("{"):
        start, end = text.find("{"), text.rfind("}")
        if start == -1 or end <= start:
            raise ValueError("No JSON object found in model response")
        text = text[start:end + 1]
    return TranscriptData.model_validate_json(text)


def extract_transcript(
    client: anthropic.Anthropic,
    path: Path,
    cache_dir: Path,
    model: str = DEFAULT_MODEL,
) -> TranscriptData:
    """Extract structured transcript data from one file, with disk caching."""
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = cache_dir / f"extract-{_file_hash(path)}.json"
    if cache_file.exists():
        log.info("Using cached extraction for %s", path.name)
        return TranscriptData.model_validate_json(cache_file.read_text(encoding="utf-8"))

    log.info("Extracting %s with %s ...", path.name, model)
    # JSON is requested via the prompt and validated client-side rather than
    # through strict structured outputs: the schema (many nullable fields)
    # exceeds the API's grammar-compilation limits ("Schema is too complex").
    messages = [{"role": "user", "content": build_content(path)}]
    data: TranscriptData | None = None
    last_error: Exception | None = None
    for attempt in range(2):
        with client.messages.stream(
            model=model,
            max_tokens=32000,
            messages=messages,
        ) as stream:
            response = stream.get_final_message()

        if response.stop_reason == "refusal":
            raise RuntimeError(f"Model declined to process {path.name}")
        if response.stop_reason == "max_tokens":
            raise RuntimeError(
                f"Extraction output for {path.name} was truncated (max_tokens); "
                "the transcript may be too long for a single pass."
            )

        raw = "".join(b.text for b in response.content if b.type == "text")
        try:
            data = _parse_json_response(raw)
            break
        except (ValueError, pydantic.ValidationError) as exc:
            last_error = exc
            log.warning("JSON parse failed for %s (attempt %d): %s",
                        path.name, attempt + 1, exc)
            # One repair pass: send the invalid output and the error back.
            messages = messages + [
                {"role": "assistant", "content": raw},
                {"role": "user", "content":
                    f"That JSON failed validation with this error:\n{exc}\n\n"
                    "Reply with ONLY the corrected JSON object."},
            ]

    if data is None:
        raise RuntimeError(f"Failed to parse extraction output for {path.name}: {last_error}")

    cache_file.write_text(data.model_dump_json(indent=2), encoding="utf-8")
    log.info(
        "Extracted %s: %d courses, %d degrees, %d transfer institutions",
        path.name, len(data.courses), len(data.degrees), len(data.transfer_institutions),
    )
    return data
